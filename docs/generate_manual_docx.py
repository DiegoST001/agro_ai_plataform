import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "configuraciones")
OUT_PATH = os.path.abspath(os.path.join(OUT_DIR, "Manual_Agro_AI_Platform.docx"))

def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text):
    doc.add_paragraph(text)

def add_list(doc, items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

def add_numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style='List Number')

def add_code_block(doc, code):
    p = doc.add_paragraph()
    for line in code.strip("\n").split("\n"):
        r = p.add_run(line + "\n")
        r.font.name = 'Consolas'
        r.font.size = Pt(9)

def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()

    # Título
    add_title(doc, "Manual del Sistema – Agro AI Platform")

    # 1. Descripción general
    add_heading(doc, "1. Descripción general", 1)
    add_paragraph(doc, "Plataforma de gestión agrícola basada en Django REST Framework.")
    add_list(doc, [
        "Módulos: usuarios, parcelas, nodos maestros/secundarios, planes, sensores, tareas, reportes, recomendaciones, alertas.",
        "Lecturas IoT en MongoDB; datos transaccionales en SQL.",
        "API documentada con drf-spectacular (OpenAPI/Swagger).",
        "Control de acceso por roles, módulos y operaciones.",
        "Integración con IA (Ollama local, Gemini, Anthropic).",
    ])

    # 2. Requisitos
    add_heading(doc, "2. Requisitos", 1)
    add_list(doc, [
        "Windows 10/11 con PowerShell.",
        "Python 3.11+ y Git.",
        "MongoDB (local o Atlas).",
        "PostgreSQL (producción) o SQLite (desarrollo).",
        "Docker Desktop (opcional).",
        "Cliente API (Postman/Insomnia).",
        "Opcional IA: Ollama local o cuenta/API key de Gemini/Anthropic.",
    ])

    # 3. Estructura básica
    add_heading(doc, "3. Estructura básica del proyecto", 1)
    add_list(doc, [
        "apps: users, parcels (parcelas), nodes, plans, ai, etc.",
        "ai/services.py: orquestación multi-proveedor de IA.",
        "nodes/views.py: ingesta y CRUD de nodos.",
        "parcels/views.py: CRUD de parcelas con permisos.",
        "users/models.py: Rol, Modulo, Operacion, RolesOperaciones, UserOperacionOverride.",
    ])

    # 4. .env
    add_heading(doc, "4. Configuración de entorno (.env)", 1)
    add_paragraph(doc, "Crea una carpeta “configuraciones” y un archivo .env en la raíz del proyecto.")
    add_code_block(doc, """
# Django
DJANGO_SECRET_KEY=pon_una_clave_segura
DJANGO_DEBUG=True
DJANGO_ALLOWED_HOSTS=127.0.0.1,localhost

# Base de datos (usa SQLite si no defines DATABASE_URL)
#DATABASE_URL=postgres://USER:PASS@HOST:5432/DBNAME

# MongoDB
MONGODB_URI=mongodb://localhost:27017
MONGODB_DB=agro_ai_platform

# IA
AI_PROVIDER=ollama          # ollama | gemini | anthropic
AI_ENDPOINT=http://localhost:11434
AI_API_KEY=                 # requerido para gemini/anthropic
    """.strip("\n"))

    add_paragraph(doc, "Notas:")
    add_list(doc, [
        "Para Mongo Atlas, usa la cadena de conexión de Atlas.",
        "Para Postgres, habilita DATABASE_URL y configura en settings.",
    ])

    # 5. Instalación local
    add_heading(doc, "5. Instalación local (Windows)", 1)
    add_numbered(doc, [
        "Clonar e instalar dependencias:",
        "Crear y activar entorno virtual.",
        "Instalar requirements, migrar y crear superusuario.",
        "Ejecutar servidor y revisar Swagger.",
    ])
    add_code_block(doc, """
git clone <repo-url>
cd agro_ai_platform
py -3.11 -m venv .venv
.\\.venv\\Scripts\\Activate.ps1
pip install -r requirements.txt
python manage.py migrate
python manage.py createsuperuser
python manage.py runserver
    """.strip("\n"))

    # 6. Modelo de roles y permisos
    add_heading(doc, "6. Modelo de roles y permisos", 1)
    add_list(doc, [
        "Rol: superadmin, administrador, técnico, agricultor.",
        "Modulo: parcelas, nodos, usuarios, etc.",
        "Operacion: ver, crear, actualizar, eliminar, aprobar, etc.",
        "RolesOperaciones: permisos por rol y módulo.",
        "UserOperacionOverride: sobrescribe permisos por usuario.",
    ])
    add_paragraph(doc, "Función central: tiene_permiso(user, modulo, accion).")
    add_paragraph(doc, "Clase DRF: HasOperationPermission('modulo','accion').")
    add_list(doc, [
        "superadmin: acceso total.",
        "administrador: acceso completo a módulos de operación.",
        "agricultor: acceso a sus propios recursos.",
    ])

    # 7. Seed de permisos (MATRIZ)
    add_heading(doc, "7. Seed de roles y permisos (MATRIZ)", 1)
    add_paragraph(doc, "Ejecuta en shell de Django para poblar roles, módulos, operaciones y permisos.")
    add_code_block(doc, """
from users.models import Rol, Modulo, Operacion, RolesOperaciones

MATRIZ = {
    'superadmin': {},
    'administrador': {
        'usuarios': ['ver','crear','actualizar','eliminar'],
        'parcelas': ['ver','crear','actualizar','eliminar'],
        'planes': ['ver','crear','actualizar','eliminar'],
        'sensores': ['ver','crear','actualizar','eliminar'],
        'recomendaciones': ['ver','aprobar'],
        'tareas': ['ver','crear','actualizar','eliminar'],
        'alertas': ['ver','actualizar'],
        'nodos': ['ver','crear','actualizar','eliminar'],
        'reportes': ['ver','crear'],
        'administracion': ['ver','actualizar'],
        'tecnicos': ['ver','crear','actualizar','eliminar'],
    },
    'tecnico': {
        'sensores': ['ver','actualizar'],
        'parcelas': ['ver'],
        'tareas': ['ver','actualizar'],
        'alertas': ['ver','actualizar'],
        'reportes': ['ver'],
        'tecnicos': ['ver','actualizar'],
    },
    'agricultor': {
        'parcelas': ['ver','crear','actualizar'],
        'planes': ['ver','crear'],
        'sensores': ['ver'],
        'tareas': ['ver','crear','actualizar'],
        'alertas': ['ver'],
        'reportes': ['ver'],
        'nodos': ['ver'],
        'recomendaciones': ['ver'],
    },
}

for rol in MATRIZ.keys():
    Rol.objects.get_or_create(nombre=rol)

MODULOS = sorted({m for perms in MATRIZ.values() for m in perms.keys()})
for m in MODULOS:
    mod, _ = Modulo.objects.get_or_create(nombre=m)
    ops = set()
    for tabla in MATRIZ.values():
        if m in tabla:
            ops |= set(tabla[m])
    for op in ops:
        Operacion.objects.get_or_create(modulo=mod, nombre=op)

for rol_nombre, tabla in MATRIZ.items():
    rol = Rol.objects.get(nombre=rol_nombre)
    if rol_nombre == 'superadmin':
        for mod in Modulo.objects.all():
            for op in mod.operaciones.all():
                RolesOperaciones.objects.get_or_create(rol=rol, modulo=mod, operacion=op)
        continue
    for m, ops in tabla.items():
        mod = Modulo.objects.get(nombre=m)
        for op_name in ops:
            op = Operacion.objects.get(modulo=mod, nombre=op_name)
            RolesOperaciones.objects.get_or_create(rol=rol, modulo=mod, operacion=op)
print("Permisos seed completado")
    """.strip("\n"))

    # 8. Endpoints clave
    add_heading(doc, "8. Endpoints y control de acceso", 1)
    add_paragraph(doc, "Parcelas (controlado por roles/permisos, sin rutas admin):")
    add_list(doc, [
        "GET/POST /parcelas/",
        "GET/PUT/PATCH/DELETE /parcelas/{id}/",
        "Admin ve todo; agricultor solo lo propio.",
    ])
    add_paragraph(doc, "Nodos:")
    add_list(doc, [
        "POST /nodes/ingest/ (ingesta maestro→Mongo, actualiza last_seen/estado).",
        "CRUD de nodos maestros/secundarios (filtrado por permisos y propiedad).",
    ])
    add_paragraph(doc, "Swagger en /api/schema/swagger/.")

    # 9. Integración IA
    add_heading(doc, "9. Integración de IA (ai/services.py)", 1)
    add_list(doc, [
        "get_active_ai_config(): obtiene configuración activa.",
        "chat_with_ai(prompt, context=None): enruta a proveedor (ollama|gemini|anthropic).",
        "ollama_chat: integración HTTP con endpoint local.",
        "get_context(usuario_id, parcela_id, consulta, prompt=None): contexto de lecturas (dia/semana/todas).",
    ])

    # 10. Docker
    add_heading(doc, "10. Despliegue con Docker (opcional)", 1)
    add_code_block(doc, """
docker build -t agro-ai .
docker run -d -p 8000:8000 --env-file .env agro-ai
    """.strip("\n"))
    add_list(doc, [
        "Producción: usar PostgreSQL (RDS) y Mongo Atlas.",
        "Configurar DJANGO_DEBUG=False y ALLOWED_HOSTS.",
    ])

    # 11. Pruebas rápidas
    add_heading(doc, "11. Pruebas rápidas de API", 1)
    add_list(doc, [
        "Autenticación: sesión admin o token.",
        "Parcelas: GET/POST /parcelas/ y GET/PUT/PATCH/DELETE /parcelas/{id}/.",
        "Ingesta IoT: POST /nodes/ingest/.",
        "Documentación: /api/schema/swagger/.",
    ])

    # 12. Troubleshooting
    add_heading(doc, "12. Solución de problemas", 1)
    add_list(doc, [
        "ImportError tiene_permiso: importa desde el módulo correcto (p.ej. parcels.permissions).",
        "403 Forbidden: verifica HasOperationPermission y que la MATRIZ esté seedada.",
        "Mongo: valida MONGODB_URI/MONGODB_DB.",
        "IA: Ollama corriendo o API key válida para Gemini/Anthropic.",
    ])

    # 13. Comandos útiles
    add_heading(doc, "13. Comandos útiles (PowerShell)", 1)
    add_code_block(doc, """
.\\.venv\\Scripts\\Activate.ps1
python manage.py makemigrations && python manage.py migrate
python manage.py createsuperuser
python manage.py runserver
py -m pip install python-docx
    """.strip("\n"))

    # 14. Buenas prácticas
    add_heading(doc, "14. Buenas prácticas", 1)
    add_list(doc, [
        "Un solo set de rutas por recurso; control por roles/permisos.",
        "Aplica permisos en get_permissions y filtra en get_queryset.",
        "Centraliza la lógica en tiene_permiso y reutilízala en todas las apps.",
        "Valida entradas y documenta con drf-spectacular.",
        "Agrega tests de permisos y propiedad de recursos.",
    ])

    doc.save(OUT_PATH)
    print(f"Manual generado en: {OUT_PATH}")

if __name__ == "__main__":
    main()